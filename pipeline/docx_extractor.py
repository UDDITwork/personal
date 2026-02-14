"""
DOCX Extraction Module
Uses python-docx to extract text, images, and tables from Word documents
"""
import base64
import time
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from loguru import logger

from .models import ExtractedImage, ExtractedTable
from config import get_session_output_dir


class DOCXExtractor:
    """Handles DOCX extraction using python-docx"""

    def __init__(self, user_id: str, session_id: str):
        self.user_id = user_id
        self.session_id = session_id
        self.output_dir = get_session_output_dir(user_id, session_id)

    def extract_docx(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract all content from DOCX file

        Args:
            docx_path: Path to the DOCX file

        Returns:
            Dictionary containing extracted text, images, tables, and metadata
        """
        logger.info(f"Starting DOCX extraction for {docx_path}")
        start_time = time.time()

        try:
            doc = Document(docx_path)

            # Extract paragraphs with style information
            paragraphs = self._extract_paragraphs(doc)

            # Extract images
            images = self._extract_images(doc)

            # Extract tables
            tables = self._extract_tables(doc)

            # Build markdown text from paragraphs
            text_markdown = self._paragraphs_to_markdown(paragraphs)

            # Build plain text
            text_plain = "\n".join([p["text"] for p in paragraphs if p["text"].strip()])

            processing_time = time.time() - start_time

            logger.success(f"DOCX extraction completed in {processing_time:.2f}s")
            logger.info(f"Extracted: {len(paragraphs)} paragraphs, {len(images)} images, {len(tables)} tables")

            return {
                "text_markdown": text_markdown,
                "text_plain": text_plain,
                "images": images,
                "tables": tables,
                "paragraphs": paragraphs,
                "processing_time": processing_time,
                "extraction_method": "python-docx"
            }

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    def _extract_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract all paragraphs with style and formatting info"""
        paragraphs = []

        for para in doc.paragraphs:
            # Skip empty paragraphs
            if not para.text.strip():
                continue

            # Determine if text is bold (check all runs)
            is_bold = any(run.bold for run in para.runs if run.bold is not None)

            # Determine if text is italic
            is_italic = any(run.italic for run in para.runs if run.italic is not None)

            paragraph_data = {
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "bold": is_bold,
                "italic": is_italic,
                "level": self._get_heading_level(para.style.name if para.style else "Normal")
            }

            paragraphs.append(paragraph_data)

        logger.debug(f"Extracted {len(paragraphs)} paragraphs")
        return paragraphs

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        if "Heading" in style_name:
            try:
                # Extract number from "Heading 1", "Heading 2", etc.
                level = int(style_name.split()[-1])
                return level
            except (ValueError, IndexError):
                return 0
        return 0

    def _extract_images(self, doc: Document) -> List[ExtractedImage]:
        """Extract all images from DOCX file"""
        images: List[ExtractedImage] = []
        image_counter = 0

        try:
            # Extract images from relationships
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_data = rel.target_part.blob
                        img_ext = rel.target_ref.split(".")[-1] if "." in rel.target_ref else "png"

                        # Save image to disk
                        image_counter += 1
                        img_filename = f"docx_img_{image_counter}.{img_ext}"
                        img_path = self.output_dir / img_filename

                        with open(img_path, "wb") as f:
                            f.write(image_data)

                        # Convert to base64
                        image_base64 = base64.b64encode(image_data).decode("utf-8")

                        # Create ExtractedImage object
                        extracted_image = ExtractedImage(
                            image_id=f"docx_img_{image_counter}",
                            page_number=0,  # DOCX doesn't have page numbers in the same way
                            image_path=str(img_path),
                            image_base64=image_base64
                        )

                        images.append(extracted_image)
                        logger.debug(f"Extracted image: {img_filename}")

                    except Exception as e:
                        logger.warning(f"Failed to extract image: {e}")

            logger.info(f"Extracted {len(images)} images from DOCX")

        except Exception as e:
            logger.error(f"Error extracting images from DOCX: {e}")

        return images

    def _extract_tables(self, doc: Document) -> List[ExtractedTable]:
        """Extract all tables from DOCX file"""
        tables: List[ExtractedTable] = []

        for table_idx, table in enumerate(doc.tables):
            try:
                # Extract all rows
                rows = []
                headers = []

                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]

                    # First row is often headers
                    if row_idx == 0:
                        headers = cells

                    rows.append(cells)

                # Build HTML representation
                html_content = self._table_to_html(rows, headers if headers else None)

                # Create ExtractedTable object
                extracted_table = ExtractedTable(
                    table_id=f"docx_table_{table_idx + 1}",
                    page_number=0,  # DOCX doesn't have page numbers
                    html_content=html_content,
                    headers=headers if headers else [],
                    rows=rows[1:] if headers else rows,  # Skip header row if detected
                    num_rows=len(rows),
                    num_cols=len(rows[0]) if rows else 0
                )

                tables.append(extracted_table)
                logger.debug(f"Extracted table {table_idx + 1} with {len(rows)} rows")

            except Exception as e:
                logger.warning(f"Failed to extract table {table_idx + 1}: {e}")

        logger.info(f"Extracted {len(tables)} tables from DOCX")
        return tables

    def _table_to_html(self, rows: List[List[str]], headers: List[str] = None) -> str:
        """Convert table data to HTML"""
        html = '<table border="1" style="border-collapse: collapse; width: 100%;">\n'

        # Add headers if provided
        if headers:
            html += "  <thead>\n    <tr>\n"
            for header in headers:
                html += f"      <th style='padding: 8px; background-color: #f2f2f2;'>{header}</th>\n"
            html += "    </tr>\n  </thead>\n"
            # Skip first row since it's the header
            data_rows = rows[1:]
        else:
            data_rows = rows

        # Add data rows
        html += "  <tbody>\n"
        for row in data_rows:
            html += "    <tr>\n"
            for cell in row:
                html += f"      <td style='padding: 8px; border: 1px solid #ddd;'>{cell}</td>\n"
            html += "    </tr>\n"
        html += "  </tbody>\n</table>"

        return html

    def _paragraphs_to_markdown(self, paragraphs: List[Dict[str, Any]]) -> str:
        """Convert paragraphs with style info to markdown"""
        markdown_lines = []

        for para in paragraphs:
            text = para["text"]
            level = para["level"]

            # Handle headings
            if level > 0:
                markdown_lines.append(f"{'#' * level} {text}\n")
            # Handle bold text
            elif para["bold"]:
                markdown_lines.append(f"**{text}**\n")
            # Handle italic text
            elif para["italic"]:
                markdown_lines.append(f"*{text}*\n")
            # Normal text
            else:
                markdown_lines.append(f"{text}\n")

        return "\n".join(markdown_lines)


async def extract_docx_complete(docx_path: str, user_id: str, session_id: str) -> Dict[str, Any]:
    """
    Convenience async function to extract complete DOCX data

    Args:
        docx_path: Path to DOCX file
        user_id: User identifier
        session_id: Session identifier

    Returns:
        Complete extraction result dictionary
    """
    extractor = DOCXExtractor(user_id, session_id)

    # Run synchronous extraction in thread pool to avoid blocking
    import asyncio
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, extractor.extract_docx, docx_path)

    return result
